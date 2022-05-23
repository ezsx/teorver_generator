import docx
import arr_generate
import main_logic
import random
def F(tasks, v , path, uniq_var): # массив заданий # количество вариантов # путь # доступныее варианты
    doc = docx.Document()
    answers = []
    variants = [x for x in range(1, v+1)] # 1-10
    print(variants)
    print(tasks)
    for j in variants: # сколько вариантов генерим
        doc.add_paragraph("                                                                            Вариант: "+str(j))
        ans=[]
        for i in tasks: # какие задания будут в варианте
            (task, answer) = main_logic.F('t'+str(i), random.choice(uniq_var.split(',')))
            doc.add_paragraph(task)
            ans.append(answer)
        answers.append(ans)
        doc.paragraphs[-1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)    # разрыв страницы
    doc.save(path+'/variants.docx')

    doc_answer = docx.Document()
    for j in range(0,len(variants)):
        doc_answer.add_paragraph("                                                                            Вариант: " + str(j+1))
        for i in range(0,len(tasks)):
            print(answers)
            print(len(answers))
            print(j-1,i-1)
            doc_answer.add_paragraph(str(str(tasks[i]) + ") " + answers[j][i]))
        doc_answer.paragraphs[-1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)  # разрыв страницы
    doc_answer.save(path+'/variants_answers.docx')